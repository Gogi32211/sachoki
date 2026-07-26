"""tz5yr_docs.py — generate the deliverable: master report + 25 per-signal .docx +
rule-database CSVs + README, then everything is zipped by the caller. ANALYSIS ONLY."""
import pandas as pd, os, shutil
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
D="/tmp/tz5yr/data/"; OUT="/tmp/tz5yr/_TZ ANALYTICS 5YR";
if os.path.exists(OUT): shutil.rmtree(OUT)
os.makedirs(OUT)
TS=['T1','T1G','T2','T2G','T3','T4','T5','T6','T9','T10','T11','T12']
ZS=['Z1','Z1G','Z2','Z2G','Z3','Z4','Z5','Z6','Z7','Z9','Z10','Z11','Z12']
L=lambda f: pd.read_csv(D+f)
bs=L('baseline_signal.csv'); by=L('baseline_year.csv')
comp=L('composites.csv'); seq=L('sequences.csv'); creg=L('composite_regime.csv'); sreg=L('sequence_regime.csv')
dim={n:L(f'dim_{n}.csv') for n in ['vol','suffix','line5','aio','prev1','vrvx','psar','rsi2']}
uni={n:L(f'universal_{n}.csv') for n in ['vol','suffix','line5','aio','prev1']}
ucomp=L('universal_composites.csv'); compy=L('composites_year.csv')
cregmap=dict(zip(creg.composite,creg.flag)); sregmap=dict(zip(sreg.signal+'|'+sreg.seq3,sreg.flag))

def H(doc,t,lvl=1): doc.add_heading(t,level=lvl)
def P(doc,t,b=False,sz=10):
    if not isinstance(b,bool): sz=b; b=False
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bool(b); r.font.size=Pt(sz); return p
def TBL(doc,df,cols,hdr=None,maxr=18):
    df=df.head(maxr); hdr=hdr or cols
    t=doc.add_table(rows=1,cols=len(cols)); t.style='Light Grid Accent 1'
    for i,h in enumerate(hdr):
        c=t.rows[0].cells[i].paragraphs[0].add_run(str(h)); c.bold=True; c.font.size=Pt(8)
    for _,row in df.iterrows():
        cells=t.add_row().cells
        for i,cc in enumerate(cols):
            v=row[cc]; v=f'{v:.3f}' if isinstance(v,float) else str(v)
            r=cells[i].paragraphs[0].add_run(v); r.font.size=Pt(8)
    doc.add_paragraph()

# ============== MASTER REPORT ==============
doc=Document()
ti=doc.add_heading('TZ + WLNBB + WICK — 5-Year Cross-Universe Deep Research',0)
P(doc,'Full historical replication & extension of the v260506 study',True,12)
P(doc,'Universe: SP500 + NASDAQ + Russell2k  ·  Period: 2021-05-26 — 2026-06-09  ·  8,073,913 analyzable bars',sz=10)
P(doc,'Dimensions: signal (T1–T12 / Z1–Z12), Line5 (L5/L12/L25/L34/L46), suffix (ED/EU/EDP/NB/NH/EB/EUR/…), '
      'A/I/O subdivision, volume bucket (W/L/N/B/VB), prev1 context, VR/VX/PSAR/RSI2, composites, 4-bar sequences.',sz=9)

H(doc,'1. Executive Summary — what changes over the full 5 years')
for b in [
 'THE HEADLINE: the v5 finding "VB volume is king (+1.92%)" REVERSES over the full history. Across all 5 years VB is the WORST bucket (cross-universe median -0.356%); B (controlled volume) is the BEST (+0.090%). The v5 VB edge was a 2025 risk-on artifact.',
 'REGIME FRAGILITY: of composites with adequate sample, 75% (422/560) are 2025-ARTIFACT (positive only in the recent regime); only 61 are STABLE across >=60% of years. For 4-bar sequences, ~57% are 2025-ARTIFACT.',
 'WHAT SURVIVES: T12L5EDP is the single most robust composite — universal (positive in all 3 universes, min +0.502%) AND STABLE (positive 4/6 years). EDP & ED suffixes, L5 line, B volume, and T9 prev1 context are the durable edges.',
 'RUSSELL2K: every raw signal is net-NEGATIVE over the full history (T1 -0.36%, T2 -0.47%) — micro/small T/Z signals carry no standalone edge once 2022 is included.',
 'CONFIRMED FROM v5: EUR suffix universally weak (-0.264%) -> deprecate. NH suffix weak (-0.204%). SP500 >> NASDAQ >> Russell2k baseline ordering holds.',
 'REVERSED FROM v5: VB (was best, now worst), and most "winning composites/sequences" do not survive 2021/2022/2024.']:
    P(doc,'•  '+b,sz=9)

H(doc,'2. Status Overview — baseline median 10d per signal, per universe')
piv=bs.pivot_table(index='signal',columns='universe',values='med10').reindex(TS+ZS).reset_index()
for u in ['sp500','nasdaq','russell2k']:
    if u not in piv: piv[u]=None
TBL(doc,piv,['signal','sp500','nasdaq','russell2k'],maxr=30)
P(doc,'Read: SP500 mostly positive; NASDAQ mixed; Russell2k uniformly negative. Strongest SP500 raw: T5 +0.59, T12 +0.58, Z6 +0.52, Z2G +0.52.',sz=9)

H(doc,'3. Cross-Universe Universal Patterns')
H(doc,'3.1 Volume bucket (THE reversal vs v5)',2)
TBL(doc,uni['vol'],['dim','cross_med','tot_n','avg_win'],['bucket','cross med10','total n','avg win%'])
P(doc,'B (controlled) is the only reliably positive bucket. VB (climactic) is the worst over the full cycle — opposite of the v5 short-window conclusion.',sz=9)
H(doc,'3.2 Suffix universality',2)
TBL(doc,uni['suffix'],['dim','cross_med','tot_n','avg_win'],['suffix','cross med10','total n','avg win%'])
H(doc,'3.3 Line5 universality',2)
TBL(doc,uni['line5'],['dim','cross_med','tot_n','avg_win'],['line5','cross med10','total n','avg win%'])
H(doc,'3.4 Prev1 context universality',2)
TBL(doc,uni['prev1'],['dim','cross_med','tot_n','avg_win'],['prev1','cross med10','total n','avg win%'])
H(doc,'3.5 A/I/O subdivision (the v5 "next step")',2)
TBL(doc,uni['aio'],['dim','cross_med','tot_n','avg_win'],['A/I/O','cross med10','total n','avg win%'])
P(doc,'O (close at/above prev body) is the least-bad subdivision overall; A and I underperform across the full history.',sz=9)

H(doc,'4. Composites — universal winners + regime stability')
H(doc,'4.1 Universal composites (positive in ALL 3 universes), by worst-universe median',2)
TBL(doc,ucomp,['composite','sp500','nasdaq','russell2k','min3','n_min'],
    ['composite','SP','NQ','R2k','min3','min n'],maxr=20)
H(doc,'4.2 Regime stability — STABLE composites (positive across >=60% of years)',2)
st=creg[creg.flag=='STABLE'].sort_values('med_all',ascending=False)
TBL(doc,st,['composite','med_all','pos_yr','tot_yr','n'],['composite','med10','+yrs','tot yrs','n'],maxr=20)
P(doc,f'STABLE: {len(creg[creg.flag=="STABLE"])}  ·  MIXED: {len(creg[creg.flag=="MIXED"])}  ·  '
      f'2025-ARTIFACT: {len(creg[creg.flag=="2025-ARTIFACT"])}.  Most high-median composites are artifacts — '
      'the STABLE set is the only deployable list.',True,9)

H(doc,'5. 4-Bar Sequences — winners by regime')
H(doc,'5.1 STABLE sequences (positive across >=60% of years)',2)
ss=sreg[sreg.flag=='STABLE'].sort_values('med_all',ascending=False)
TBL(doc,ss,['signal','seq3','med_all','pos_yr','tot_yr','n'],['signal','seq [-3|-2|-1]','med10','+yrs','tot','n'],maxr=22)
P(doc,f'STABLE: {len(sreg[sreg.flag=="STABLE"])}  ·  MIXED: {len(sreg[sreg.flag=="MIXED"])}  ·  '
      f'2025-ARTIFACT: {len(sreg[sreg.flag=="2025-ARTIFACT"])}.  The very-high-median sequences (+6–8%) are almost all '
      'single-year artifacts; the STABLE ones cluster around +2–5% and recur (Z1G as a building block dominates).',9)

H(doc,'6. VR / PSAR / RSI2 effects (cross-universe avg median10)')
for nm,lab in [('vrvx','VR / VX (VIX state)'),('psar','PSAR (PS / PB)'),('rsi2','RSI2 (R2L/H/X/D)')]:
    g=dim[nm].groupby('dim').agg(cross=('med10','mean'),n=('n','sum')).reset_index().sort_values('cross',ascending=False)
    H(doc,'6.x '+lab,3); TBL(doc,g,['dim','cross','n'],[lab,'cross med10','total n'],maxr=8)

H(doc,'7. Conclusions')
for c in [
 '1. Regime is everything. The single biggest correction to v5 is that VB-volume and most composites/sequences are 2025-regime-dependent. Always validate per-year before deploying.',
 '2. The durable core: B-volume + L5 + EDP/ED suffix + T9/Z5 prev1, with T12L5EDP as the flagship universal+stable composite.',
 '3. Russell2k needs different treatment: raw signals net-negative; only specific STABLE composites/sequences (and mean-reversion context) work, with small size.',
 '4. Deprecate EUR and de-rate NH and A/I "A" suffix. Prefer O subdivision and EDP/ED.',
 '5. Sequence + composite together still beats either alone, but only the STABLE subset should drive scanner rules.',
 '6. The per-signal deep-dive files and rule_database_5yr.csv give the full backing tables; regime flags are attached to every composite and sequence.']:
    P(doc,c,sz=9)
doc.save(OUT+'/_TZ_WLNBB_Research_Report_5YR.docx')
print('master report saved')

# ============== PER-SIGNAL DOCS ==============
for sg in TS+ZS:
    d=Document(); d.add_heading(f'{sg} — 5-Year Signal Analytics',0)
    P(d,f'TZ+WLNBB+WICK  ·  SP500 + NASDAQ + Russell2k  ·  2021-05 — 2026-06  ·  full history',sz=9)
    b=bs[bs.signal==sg]
    H(d,'A. Baseline per universe (fwd_10d)')
    TBL(d,b,['universe','n','med10','avg10','med5','med20','win','fail','big_win'],
        ['univ','n','med10','avg10','med5','med20','win%','fail%','bigwin%'])
    H(d,'B. Per-year regime (median 10d, win%)')
    yy=by[by.signal==sg].pivot_table(index='universe',columns='yr',values='med10').reset_index()
    TBL(d,yy,['universe']+[c for c in yy.columns if c!='universe'])
    for nm,lab in [('vol','C. Volume bucket'),('line5','D. Line5'),('suffix','E. Suffix'),('aio','F. A/I/O'),('prev1','G. Prev1 context')]:
        sub=dim[nm][dim[nm].signal==sg].sort_values('med10',ascending=False)
        if len(sub): H(d,lab); TBL(d,sub,['universe','dim','n','med10','win','fail','big_win'],
            ['univ',lab.split('. ')[1],'n','med10','win%','fail%','bigwin%'],maxr=14)
    cc=comp[comp.signal==sg].sort_values('med10',ascending=False)
    if len(cc):
        cc=cc.copy(); cc['regime']=cc.composite.map(cregmap).fillna('—')
        H(d,'H. Top composites (signal × Line5 × suffix)')
        TBL(d,cc,['composite','universe','n','med10','win','big_win','regime'],
            ['composite','univ','n','med10','win%','bigwin%','regime'],maxr=18)
    sq=seq[seq.signal==sg].sort_values('med10',ascending=False)
    if len(sq):
        sq=sq.copy(); sq['regime']=(sq.signal+'|'+sq.seq3).map(sregmap).fillna('—')
        H(d,'I. Top 4-bar sequences  [bar-3 | bar-2 | bar-1] → '+sg)
        TBL(d,sq,['seq3','universe','n','med10','win','regime'],
            ['seq [-3|-2|-1]','univ','n','med10','win%','regime'],maxr=18)
    d.save(f'{OUT}/{sg}_analytics.docx')
print('per-signal docs saved (%d)'%len(TS+ZS))

# ============== CSVs + README ==============
# rule database = composites + sequences with regime
cdb=comp.copy(); cdb['regime']=cdb.composite.map(cregmap).fillna('—'); cdb['type']='composite'
sdb=seq.copy(); sdb['regime']=(sdb.signal+'|'+sdb.seq3).map(sregmap).fillna('—'); sdb['type']='sequence'
cdb.to_csv(OUT+'/rule_database_composites_5yr.csv',index=False)
sdb.to_csv(OUT+'/rule_database_sequences_5yr.csv',index=False)
by.to_csv(OUT+'/baseline_per_year.csv',index=False)
creg.to_csv(OUT+'/composite_regime_flags.csv',index=False); sreg.to_csv(OUT+'/sequence_regime_flags.csv',index=False)
for n in ['vol','suffix','line5','aio','prev1']: uni[n].to_csv(OUT+f'/universal_{n}.csv',index=False)
with open(OUT+'/README.md','w') as f:
    f.write('''# TZ + WLNBB + WICK — 5-Year Research (full replication of v260506)

**Source:** studio_analytics.duckdb · 2021-05-26 — 2026-06-09 · 8,073,913 analyzable bars · SP500 + NASDAQ + Russell2k.
Replicates every dimension of the v5 study (signal, Line5, suffix, A/I/O, volume bucket, prev1, VR/VX/PSAR/RSI2,
composites, 4-bar sequences) over the FULL history, adding a per-year regime-stability flag to every pattern.

## Files
- `_TZ_WLNBB_Research_Report_5YR.docx` — master report (exec summary, universal patterns, regime stability, conclusions).
- `<SIGNAL>_analytics.docx` — per-signal deep dive (T1–T12, Z1–Z12): baseline, per-year, volume, Line5, suffix, A/I/O, prev1, top composites, top sequences.
- `rule_database_composites_5yr.csv` / `rule_database_sequences_5yr.csv` — every pattern with n/med/win/fail/big_win + regime flag.
- `composite_regime_flags.csv` / `sequence_regime_flags.csv` — STABLE / MIXED / 2025-ARTIFACT classification.
- `baseline_per_year.csv` — per signal per year.
- `universal_*.csv` — cross-universe universality tables.

## Definitions
- med10/med5/med20 = median forward return at 10/5/20 days (%). win% = P(fwd_10d>0). fail% = P(fwd_10d<=-5%). big_win% = P(fwd_10d>=+5%).
- composite = signal + Line5 + suffix (e.g. T12L5EDP). sequence = [bar-3|bar-2|bar-1] codes ending at the signal bar.
- regime flag: STABLE = positive median in >=60% of years (>=4 yrs); 2025-ARTIFACT = positive only in <=40% of years; else MIXED.

## Headline corrections vs v5
1. VB volume REVERSES: best in the 2025 window, WORST over 5 years (-0.356 cross-universe). B is the durable bucket.
2. ~75% of composites and ~57% of sequences are 2025-ARTIFACT. Only the STABLE subset is deployable.
3. T12L5EDP is the flagship: universal + stable. EUR/NH suffix weak (confirmed). Russell2k raw signals net-negative.
''')
print('CSVs + README saved -> '+OUT)
