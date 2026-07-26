"""tz5yr_docs2.py — FULL per-signal docs matching the old reports section-for-section
(A Executive+Status, C baseline+MFE/MAE/RR+per-year, D composites+Status, E top+reject
sequences, F prev1, G L+volume, H suffix+AIO, I MA50-reclaim+price-bucket) + regime flags.
Regenerates master report too. ANALYSIS ONLY."""
import pandas as pd, os, shutil
from docx import Document
from docx.shared import Pt
D="/tmp/tz5yr/data/"; OUT="/tmp/tz5yr/_TZ ANALYTICS 5YR"
if os.path.exists(OUT): shutil.rmtree(OUT)
os.makedirs(OUT)
TS=['T1','T1G','T2','T2G','T3','T4','T5','T6','T9','T10','T11','T12']
ZS=['Z1','Z1G','Z2','Z2G','Z3','Z4','Z5','Z6','Z7','Z9','Z10','Z11','Z12']
UU=['sp500','nasdaq','russell2k']; UL={'sp500':'SP','nasdaq':'NQ','russell2k':'R2k'}
Ld=lambda f: pd.read_csv(D+f)
rb=Ld('rich_baseline.csv'); rc=Ld('rich_composite.csv'); rs=Ld('rich_sequence.csv')
rp=Ld('rich_prev1.csv'); rl=Ld('rich_line5.csv'); rv=Ld('rich_vol.csv'); rsf=Ld('rich_suffix.csv')
rai=Ld('rich_aio.csv'); rre=Ld('rich_reclaim.csv'); rpb=Ld('rich_pbucket.csv')
by=Ld('baseline_year.csv'); creg=Ld('composite_regime.csv'); sreg=Ld('sequence_regime.csv')
cregmap=dict(zip(creg.composite,creg.flag)); sregmap=dict(zip(sreg.signal+'|'+sreg.seq3,sreg.flag))
uni={n:Ld(f'universal_{n}.csv') for n in ['vol','suffix','line5','aio','prev1']}
ucomp=Ld('universal_composites.csv'); dimv=Ld('dim_vrvx.csv');dimp=Ld('dim_psar.csv');dimr=Ld('dim_rsi2.csv')
# NEW dims (engine3): suffix components + subdivided suffix + body/wick + gap/range
rne=Ld('rich_ne.csv'); rwk=Ld('rich_wickc.csv'); rpn=Ld('rich_penc.csv'); rcl=Ld('rich_cls.csv')
rsfa=Ld('rich_sfxaio.csv'); rbw=Ld('rich_bodywick.csv'); rbs=Ld('rich_bodysz.csv'); rws=Ld('rich_wshape.csv'); rgr=Ld('rich_gaprange.csv')
ratm=Ld('rich_atomic.csv')
# line5 (engine5): PSAR / RSI2 / VIX / full bar_line5
rpsar=Ld('rich_psar.csv'); rrsi2=Ld('rich_rsi2.csv'); rvrvx=Ld('rich_vrvx.csv'); rl5f=Ld('rich_line5full.csv'); rpr=Ld('rich_psarrsi2.csv')

DESC={'T1':'Bullish reversal: prev bar bearish, current opens within prior body (open≥close[1], ≤open[1]), closes above prior open.',
'T1G':'Bullish gap reversal strong: prev bar bearish, current opens fully above prior body (gap up), closes above prior open.',
'T2':'Bullish continuation: prev bar bullish, current opens at/above prev open, closes above prev close.',
'T2G':'Bullish continuation strong: prev bar bullish, current opens at/above prev open, closes above prev close (stronger).',
'T3':'Bullish inside reversal: prev bar bearish, current bullish, opens inside prev body, closes inside prev range.',
'T4':'Bullish engulfing: current bullish bar fully engulfs previous bearish bar body (or wick). Strongest bullish reversal.',
'T5':'Partial bullish recovery: prev bar bearish, current bullish, opens below prev low, closes below prev open. Weak recovery.',
'T6':'Bullish continuation engulf: prev bar bullish, current fully engulfs previous bullish bar. Acceleration signal.',
'T9':'Bullish inside bar: current bullish bar range fully inside previous bar range. Compression before breakout.',
'T10':'Bullish inside continuation: prev bar bullish, current bullish fully inside previous range. Tight consolidation.',
'T11':'Bullish pullback: prev bar bullish, current opens below prev open but closes above prev close. Higher close after dip.',
'T12':'Bullish gap-down open: prev bar bullish, current opens & closes below prev open but close>open (still bullish bar).',
'Z1':'Bearish gap reversal: prev bar bullish, current bearish, opens within prev body & below prev close, closes below prev open.',
'Z1G':'Bearish gap reversal strong: prev bar bullish, current bearish, opens fully below prev body, closes below prev open.',
'Z2':'Bearish continuation: prev bar bearish, current bearish, opens within prev body, closes below prev close.',
'Z2G':'Bearish continuation strong: prev bar bearish, current opens at/below prev open, closes below prev close.',
'Z3':'Bearish inside reversal: prev bar bullish, current bearish, opens above prev range, closes inside prev range.',
'Z4':'Bearish engulfing: current bearish bar fully engulfs previous bullish bar. Strongest bearish reversal.',
'Z5':'Partial bearish reversal: prev bar bullish, current bearish, opens above prev high, closes above prev open. Weak reversal.',
'Z6':'Bearish continuation engulf: prev bar bearish, current fully engulfs previous bearish bar. Acceleration to downside.',
'Z7':'Doji: close equals open. Indecision — only shown when no other T/Z signal present.',
'Z9':'Bearish inside bar: current bearish bar range fully inside previous bar range. Compression — unresolved.',
'Z10':'Bearish inside continuation: prev bar bearish, current bearish fully inside previous range. Tight consolidation.',
'Z11':'Bearish pullback: prev bar bearish, current opens above prev open but closes below prev close. Lower close after bounce.',
'Z12':'Bearish gap-up open: prev bar bullish, current opens at/below prev open, closes below open (bearish gap-down).'}

def status(med,fail):
    if pd.isna(med): return '—'
    if med>=0.7 and fail<=20: return 'GOOD'
    if med<=-0.1 or fail>=28: return 'REJECT'
    return 'AVERAGE'
def H(d,t,l=1): d.add_heading(t,level=l)
def P(d,t,b=False,sz=9):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=bool(b); r.font.size=Pt(sz); return p
def fmt(v):
    if isinstance(v,float): return f'{v:.3f}'
    return str(v)
def TBL(d,rows,hdr,maxr=20):
    rows=rows[:maxr]
    t=d.add_table(rows=1,cols=len(hdr)); t.style='Light Grid Accent 1'
    for i,h in enumerate(hdr):
        r=t.rows[0].cells[i].paragraphs[0].add_run(str(h)); r.bold=True; r.font.size=Pt(7.5)
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            r=c[i].paragraphs[0].add_run(fmt(v)); r.font.size=Pt(7.5)
    d.add_paragraph()

def sidebyside(df, sg, dimlabel, maxr=16, sort_med=True):
    """pivot a rich-dim df to: dim | SP n/med/fail | NQ ... | R2k ..."""
    sub=df[df.signal==sg]
    dims=sub['dim'].unique() if 'dim' in sub else []
    rows=[]
    for dv in dims:
        rec=[dv]
        med_sp=None
        for u in UU:
            r=sub[(sub.dim==dv)&(sub.universe==u)]
            if len(r):
                r=r.iloc[0]; rec += [int(r['n']), r['m10'], r['fail']]
                if u=='sp500': med_sp=r['m10']
            else: rec += ['—','—','—']
        rows.append((rec, med_sp if med_sp is not None else -99))
    if sort_med: rows.sort(key=lambda x:(x[1] if isinstance(x[1],(int,float)) else -99), reverse=True)
    hdr=[dimlabel]+[f'{UL[u]} n' for u in UU for _ in [0]][:0]
    hdr=[dimlabel,'SP n','SP med','SP fail','NQ n','NQ med','NQ fail','R2k n','R2k med','R2k fail']
    return [r[0] for r in rows][:maxr], hdr

# ================= MASTER (regenerate, richer exec) =================
doc=Document(); doc.add_heading('TZ + WLNBB + WICK — 5-Year Cross-Universe Deep Research',0)
P(doc,'Full historical replication & extension of the v260506 study',True,12)
P(doc,'SP500 + NASDAQ + Russell2k · 2021-05-26 — 2026-06-09 · 8,073,913 analyzable bars',sz=10)
H(doc,'1. Executive Summary — what changes over the full 5 years')
for b in ['THE HEADLINE: v5 "VB volume is king (+1.92%)" REVERSES over the full history — VB is the WORST bucket (cross-universe median -0.356%); B (controlled) is BEST (+0.090%). v5 VB edge was a 2025 risk-on artifact.',
 'REGIME FRAGILITY: 75% of composites (422/560) and ~57% of 4-bar sequences are 2025-ARTIFACT (positive only recently). Only 61 composites are STABLE across >=60% of years.',
 'DURABLE CORE: T12L5EDP (universal + stable). EDP/ED suffix, L5 line, B volume, T9 prev1.',
 'RUSSELL2K: every raw signal net-NEGATIVE over the full history (2022 bear included).',
 'CONFIRMED v5: EUR suffix weak (deprecate), NH weak, SP500>>NASDAQ>>Russell2k.',
 'ATOMIC DECOMPOSITION (new): the suffix edge is ~80% in ONE component — close (A/I/O). On bull (T) signals, close=O (weak close below prior body) gives +0.30 lift universally; EO combo (escape+weak-close) gives +0.5-0.6. GAP is the other dominant axis: G3/G2 gap-up = +0.4-0.6 lift; no-gap is worst. See ATOMIC_SUFFIX_GAP_DISCOVERIES.md and per-signal sections H3/H4/J.']:
    P(doc,'•  '+b)
H(doc,'2. Baseline median 10d per signal × universe')
piv=rb.pivot_table(index='signal',columns='universe',values='m10').reindex(TS+ZS)
TBL(doc,[[s]+[piv.loc[s,u] if u in piv and s in piv.index else '—' for u in UU] for s in TS+ZS],['signal','SP med10','NQ med10','R2k med10'],maxr=30)
H(doc,'3. Cross-universe universal patterns')
for nm,lab in [('vol','3.1 Volume bucket (reversal vs v5)'),('suffix','3.2 Suffix'),('line5','3.3 Line5'),('prev1','3.4 Prev1'),('aio','3.5 A/I/O')]:
    H(doc,lab,2); u=uni[nm]
    TBL(doc,[[r['dim'],r['cross_med'],int(r['tot_n']),round(r['avg_win'],1)] for _,r in u.iterrows()],[lab.split(' ',1)[1],'cross med10','total n','avg win%'],maxr=14)
H(doc,'4. Universal composites (positive in all 3) + STABLE set')
TBL(doc,[[r['composite'],r['sp500'],r['nasdaq'],r['russell2k'],r['min3'],int(r['n_min'])] for _,r in ucomp.head(18).iterrows()],['composite','SP','NQ','R2k','min3','min n'])
st=creg[creg.flag=='STABLE'].sort_values('med_all',ascending=False)
P(doc,f'STABLE composites: {len(creg[creg.flag=="STABLE"])} · MIXED: {len(creg[creg.flag=="MIXED"])} · 2025-ARTIFACT: {len(creg[creg.flag=="2025-ARTIFACT"])}',True)
TBL(doc,[[r['composite'],r['med_all'],f"{r['pos_yr']}/{r['tot_yr']}",int(r['n'])] for _,r in st.head(18).iterrows()],['STABLE composite','med10','+yrs','n'])
H(doc,'5. VR / PSAR / RSI2 (cross-universe avg med10)')
for df,lab in [(dimv,'VR/VX'),(dimp,'PSAR'),(dimr,'RSI2')]:
    g=df.groupby('dim').agg(cross=('med10','mean'),n=('n','sum')).reset_index().sort_values('cross',ascending=False)
    H(doc,lab,3); TBL(doc,[[r['dim'],round(r['cross'],3),int(r['n'])] for _,r in g.iterrows()],[lab,'cross med10','n'],maxr=8)
H(doc,'6. Conclusions')
for c in ['Regime is everything — validate per-year before deploying any v5 winner.','Durable core: B-vol + L5 + EDP/ED + T9 prev1, flagship T12L5EDP.','Russell2k raw signals net-negative; only STABLE composites/sequences work, small size.','Deprecate EUR, de-rate NH and "A" subdivision; prefer O and EDP/ED.','Per-signal files + rule_database CSVs carry the full backing tables with regime flags.']:
    P(doc,c)
doc.save(OUT+'/_TZ_WLNBB_Research_Report_5YR.docx'); print('master saved')

# ================= PER-SIGNAL (FULL structure) =================
for sg in TS+ZS:
    d=Document(); d.add_heading(f'{sg} — 5-Year Signal Analytics',0)
    P(d,'TZ+WLNBB+WICK Research Report',True,11)
    P(d,DESC.get(sg,''),sz=10)
    P(d,'SP500 + NASDAQ + Russell2k · daily · 2021-05-26 – 2026-06-09 · full history (8.07M bars)',sz=9)
    b=rb[rb.signal==sg]
    # A. Executive Conclusion
    H(d,'A. Executive Conclusion')
    arows=[]
    for u in UU:
        r=b[b.universe==u]
        if len(r):
            r=r.iloc[0]; rr=round(r['mfe']/abs(r['mae']),3) if r['mae'] else '—'
            arows.append([UL[u],status(r['m10'],r['fail']),r['m10'],r['fail'],f"n={int(r['n'])} rr={rr}"])
    TBL(d,arows,['Universe','Status','Med 10D','Fail%','Notes'])
    for u in UU:
        cc=rc[(rc.signal==sg)&(rc.universe==u)&(rc.n>=40)].sort_values('m10',ascending=False)
        if len(cc): r=cc.iloc[0]; P(d,f'{UL[u]} best composite: {r["composite"]} (med {r["m10"]:.3f}%, fail {r["fail"]}%, n={int(r["n"])})')
    sq=rs[(rs.signal==sg)&(rs.n>=20)]
    if len(sq):
        bb=sq.sort_values('m10',ascending=False).iloc[0]; ww=sq.sort_values('m10').iloc[0]
        P(d,f'Best sequence: {bb["seq3"]} → {sg} (med {bb["m10"]:.3f}%, n={int(bb["n"])})')
        P(d,f'Worst sequence: {ww["seq3"]} → {sg} (med {ww["m10"]:.3f}%, fail {ww["fail"]}%, n={int(ww["n"])})')
    # C. Baseline Statistics
    H(d,'C. Baseline Statistics (per universe)')
    for u in UU:
        r=b[b.universe==u]
        if not len(r): continue
        r=r.iloc[0]; rr=round(r['mfe']/abs(r['mae']),3) if r['mae'] else '—'
        P(d,f'{UL[u]} — All {sg} (n={int(r["n"])})',True,9)
        TBL(d,[['Avg return',r['a1'],r['a3'],r['a5'],r['a10']],
               ['Median','—','—',r['m5'],r['m10']]],
            ['Metric','1D','3D','5D','10D'])
        P(d,f"Win% {r['win']} · Big win(>10%) {r['bigwin']}% · Fail(<-5%) {r['fail']}% · "
            f"Avg MFE/MAE {r['mfe']:.3f}%/{r['mae']:.3f}% · Reward/Risk {rr} · Med 20D {r['m20']}",sz=8)
    # C2. per-year regime (my addition)
    H(d,'C2. Per-year regime — median 10d')
    yy=by[by.signal==sg]
    yrs=sorted(yy.yr.unique())
    rows=[]
    for u in UU:
        row=[UL[u]]
        for y in yrs:
            c=yy[(yy.universe==u)&(yy.yr==y)]
            row.append(c.iloc[0]['med10'] if len(c) else '—')
        rows.append(row)
    TBL(d,rows,['Univ']+[str(y) for y in yrs])
    # D. Composite ranking per universe
    H(d,'D. Composite Ranking (signal × Line5 × suffix)')
    for u in UU:
        cc=rc[(rc.signal==sg)&(rc.universe==u)&(rc.n>=40)].sort_values('m10',ascending=False)
        if not len(cc): continue
        P(d,f'D-{UL[u]} — top composites',True,9)
        TBL(d,[[r['composite'],int(r['n']),r['a10'],r['m10'],r['fail'],status(r['m10'],r['fail']),cregmap.get(r['composite'],'—')] for _,r in cc.head(14).iterrows()],
            ['Composite','n','Avg10D','Med10D','Fail%','Status','Regime'])
    # E. 4-bar sequences top + reject per universe
    H(d,'E. 4-Bar Sequences  [bar-3 | bar-2 | bar-1] → '+sg)
    for u in UU:
        sq=rs[(rs.signal==sg)&(rs.universe==u)&(rs.n>=20)]
        if not len(sq): continue
        top=sq.sort_values('m10',ascending=False).head(10); rej=sq.sort_values('m10').head(6)
        P(d,f'E-{UL[u]} — top sequences (n>=20)',True,9)
        TBL(d,[[r['seq3'],int(r['n']),r['a10'],r['m10'],r['fail'],status(r['m10'],r['fail']),sregmap.get(sg+'|'+r['seq3'],'—')] for _,r in top.iterrows()],
            ['Sequence','n','Avg10D','Med10D','Fail%','Status','Regime'])
        P(d,f'E-{UL[u]} — reject sequences',True,9)
        TBL(d,[[r['seq3'],int(r['n']),r['a10'],r['m10'],r['fail'],'REJECT'] for _,r in rej.iterrows()],
            ['Sequence','n','Avg10D','Med10D','Fail%','Status'])
    # F. Prev1
    H(d,'F. Prev1 Signal Effect (sorted by SP median)')
    rows,hdr=sidebyside(rp,sg,'Prev1')
    TBL(d,rows,hdr)
    # G. L + volume
    H(d,'G. Volume / L / WLNBB Analysis')
    P(d,'G1. Line5 (L-signal) breakdown',True,9); r1,h1=sidebyside(rl,sg,'Line5'); TBL(d,r1,h1)
    P(d,'G2. Volume bucket',True,9); r2,h2=sidebyside(rv,sg,'Volume'); TBL(d,r2,h2)
    # H. suffix + aio
    H(d,'H. Wick / Suffix Analysis')
    P(d,'H1. Full suffix breakdown',True,9); r3,h3=sidebyside(rsf,sg,'Suffix'); TBL(d,r3,h3)
    P(d,'H2. A/I/O subdivision',True,9); r4,h4=sidebyside(rai,sg,'A/I/O'); TBL(d,r4,h4)
    # I. reclaim + price bucket
    H(d,'I. MA50 Reclaim + Price Bucket')
    P(d,'I1. MA50 reclaim',True,9); r5,h5=sidebyside(rre,sg,'MA50'); TBL(d,r5,h5)
    P(d,'I2. Price bucket',True,9); r6,h6=sidebyside(rpb,sg,'Price',maxr=8,sort_med=False); TBL(d,r6,h6)
    # H3. Suffix decomposition into atomic components (separate lines)
    H(d,'H3. Suffix components — analyzed SEPARATELY')
    P(d,'suffix = ne(E/N escape) + wick(B/U/D) + pen(H/P/R) + close(A/I/O). Each component, isolated:',sz=8)
    for df,lab in [(rne,'ne (escape E/N)'),(rwk,'wick (B/U/D)'),(rpn,'pen (H/P/R)'),(rcl,'close (A/I/O)')]:
        P(d,lab,True,9); rr,hh=sidebyside(df,sg,lab.split(' ')[0]); TBL(d,rr,hh,maxr=8)
    # H4. Subdivided suffix (A/I/O) — the v5 'next step'
    H(d,'H4. Subdivided suffix (full + A/I/O, e.g. EBO vs EBA)')
    rr,hh=sidebyside(rsfa,sg,'sfx+AIO'); TBL(d,rr,hh,maxr=16)
    # J. Body/Wick (line3) + Gap/Range (line4)
    H(d,'J. Bar shape — Body/Wick (line 3) & Gap/Range (line 4)')
    P(d,'J1. Body size (X/M/S)',True,9); rr,hh=sidebyside(rbs,sg,'body'); TBL(d,rr,hh,maxr=6)
    P(d,'J2. Wick shape (F/J/TB/BB)',True,9); rr,hh=sidebyside(rws,sg,'wick'); TBL(d,rr,hh,maxr=8)
    P(d,'J3. Body/wick combined',True,9); rr,hh=sidebyside(rbw,sg,'body/wick'); TBL(d,rr,hh,maxr=14)
    P(d,'J4. Gap × Range (G1/G2/G3 × N/C/V)',True,9); rr,hh=sidebyside(rgr,sg,'gap×range'); TBL(d,rr,hh,maxr=14)
    # L. Line 5 — PSAR / RSI2 / VIX (bar_line5)
    H(d,'L. Line 5 (bar_line5) — PSAR / RSI2 / VIX')
    P(d,'bar_line5 = [VR|VX VIX-state]-[PS|PB PSAR]-[R2L/H/X/D RSI2]. Key 5-yr finding: '
      'R2L (oversold RSI2) is the gem (+0.3 lift universal, the "buy-weakness" theme); PB+R2L strongest combo.',sz=8)
    P(d,'L1. PSAR (PS/PB)',True,9); rr,hh=sidebyside(rpsar,sg,'PSAR'); TBL(d,rr,hh,maxr=4)
    P(d,'L2. RSI2 (R2L/R2H/R2X/R2D)',True,9); rr,hh=sidebyside(rrsi2,sg,'RSI2'); TBL(d,rr,hh,maxr=6)
    P(d,'L3. VIX state (VR/VX)',True,9); rr,hh=sidebyside(rvrvx,sg,'VIX'); TBL(d,rr,hh,maxr=4)
    P(d,'L4. PSAR × RSI2 combo',True,9); rr,hh=sidebyside(rpr,sg,'PSAR×RSI2'); TBL(d,rr,hh,maxr=10)
    P(d,'L5. bar_line5 full',True,9); rr,hh=sidebyside(rl5f,sg,'bar_line5'); TBL(d,rr,hh,maxr=14)
    # K. Atomic profile lift on THIS signal (weak-close gap-up)
    H(d,'K. Atomic profile — close=O / gap filters on '+sg)
    P(d,'Discovery: the suffix edge is ~80% in close (A/I/O) + gap. Below: this signal filtered by '
      'close=O (weak close, bull edge), close=O+gap, EO+gap, and the bear-side close=A. '
      'NB: engulf signals (T4/T6/Z4/Z6) close strong by definition → close=O rare.',sz=8)
    order=['all','close=O','close=O+gap','EO+gap','close=A (bear-side)','A+gap']
    sub=ratm[ratm.signal==sg]
    arows=[]
    for v in order:
        rec=[v]; ok=False
        for u in UU:
            r=sub[(sub.variant==v)&(sub.universe==u)]
            if len(r): r=r.iloc[0]; rec+=[int(r['n']),r['m10'],r['win']]; ok=True
            else: rec+=['—','—','—']
        if ok: arows.append(rec)
    TBL(d,arows,['variant','SP n','SP med','SP win','NQ n','NQ med','NQ win','R2k n','R2k med','R2k win'])
    P(d,'SPLIT analytic · TZ+WLNBB+WICK 5-Year Research · regime flags: STABLE / MIXED / 2025-ARTIFACT',sz=8)
    d.save(f'{OUT}/{sg}_analytics.docx')
print('per-signal docs:',len(TS+ZS))

# CSVs + README
rc2=rc.copy(); rc2['status']=rc2.apply(lambda r:status(r['m10'],r['fail']),axis=1); rc2['regime']=rc2.composite.map(cregmap).fillna('—')
rs2=rs.copy(); rs2['status']=rs2.apply(lambda r:status(r['m10'],r['fail']),axis=1); rs2['regime']=(rs2.signal+'|'+rs2.seq3).map(sregmap).fillna('—')
rc2.to_csv(OUT+'/rule_database_composites_5yr.csv',index=False)
rs2.to_csv(OUT+'/rule_database_sequences_5yr.csv',index=False)
rb.to_csv(OUT+'/rich_baseline_all.csv',index=False); by.to_csv(OUT+'/baseline_per_year.csv',index=False)
creg.to_csv(OUT+'/composite_regime_flags.csv',index=False); sreg.to_csv(OUT+'/sequence_regime_flags.csv',index=False)
for n in ['vol','suffix','line5','aio','prev1']: uni[n].to_csv(OUT+f'/universal_{n}.csv',index=False)
shutil.copy('/Users/sachoki/Desktop/sachoki-desktop/analysis/_readme_5yr.md', OUT+'/README.md') if os.path.exists('/Users/sachoki/Desktop/sachoki-desktop/analysis/_readme_5yr.md') else None
shutil.copy('/Users/sachoki/Desktop/sachoki-desktop/SUFFIX_DECOMP_DISCOVERIES.md', OUT+'/ATOMIC_SUFFIX_GAP_DISCOVERIES.md')
if os.path.exists('/Users/sachoki/Desktop/sachoki-desktop/ATOMIC_BACKTEST.md'):
    shutil.copy('/Users/sachoki/Desktop/sachoki-desktop/ATOMIC_BACKTEST.md', OUT+'/ATOMIC_PROFILE_BACKTEST.md')
ratm.to_csv(OUT+'/atomic_profile_per_signal.csv',index=False)
if os.path.exists('/Users/sachoki/Desktop/sachoki-desktop/LINE5_DISCOVERIES.md'):
    shutil.copy('/Users/sachoki/Desktop/sachoki-desktop/LINE5_DISCOVERIES.md', OUT+'/LINE5_PSAR_RSI2_VIX_DISCOVERIES.md')
rrsi2.to_csv(OUT+'/line5_rsi2.csv',index=False); rpr.to_csv(OUT+'/line5_psar_x_rsi2.csv',index=False)
rne.to_csv(OUT+'/component_ne.csv',index=False); rcl.to_csv(OUT+'/component_close.csv',index=False)
rsfa.to_csv(OUT+'/suffix_subdivided_aio.csv',index=False); rgr.to_csv(OUT+'/gap_range.csv',index=False); rbw.to_csv(OUT+'/body_wick.csv',index=False)
print('DONE ->',OUT)
