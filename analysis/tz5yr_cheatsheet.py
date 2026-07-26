"""tz5yr_cheatsheet.py — one-page master cheat-sheet of all 5-year atomic discoveries
(bull/bear profiles, winning recipe, avoid-list, regime note). Writes .docx + .md
into the ZIP folder. ANALYSIS ONLY."""
from docx import Document
from docx.shared import Pt
import os
OUT="/tmp/tz5yr/_TZ ANALYTICS 5YR"
def H(d,t,l=1): d.add_heading(t,level=l)
def P(d,t,b=False,sz=9):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=bool(b); r.font.size=Pt(sz); return p
def TBL(d,rows,hdr):
    t=d.add_table(rows=1,cols=len(hdr)); t.style='Light Grid Accent 1'
    for i,h in enumerate(hdr):
        r=t.rows[0].cells[i].paragraphs[0].add_run(str(h)); r.bold=True; r.font.size=Pt(8)
    for row in rows:
        c=t.add_row().cells
        for i,v in enumerate(row):
            r=c[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(8)
    d.add_paragraph()

d=Document(); d.add_heading('TZ + WLNBB — 5-Year Master Cheat-Sheet',0)
P(d,'All atomic discoveries on one page · 8.07M bars · 2021-05 — 2026-06 · SP500+NASDAQ+Russell2k',True,10)
P(d,'lift = median fwd_10d minus the side baseline (T or Z), %. Verified across all 3 universes.',sz=8)

H(d,'★ THE WINNING EDGE — "weak-close gap-up" (best of the whole study)')
P(d,'Rule: BULL T-signal AND close=O (weak close, below prior body) AND gap∈(G2,G3). Entry next-open, '
  '−15% stop / +100% target, $500k entry-$vol floor, small fractional size.',True,9)
TBL(d,[['sp500','+0.84','18,996','55%'],['russell2k (EO+gap)','+0.70','44,009','53%'],
       ['nasdaq (EO+gap)','+0.51','23,133','51%']],
    ['universe','EXPECT %/trade','n','win%'])
P(d,'Per-year (nas+r2k): 2021 +0.23 · 2022 −0.45 · 2023 +0.28 · 2024 +1.15 · 2025 +0.88 · 2026 +0.84 '
  '→ POSITIVE 5/6 years, only the 2022 bear negative. Stand down in confirmed RISK_OFF.',True,9)

H(d,'BULL profile — atomic components (T-signals), lift med10 [SP / NQ / R2k]')
TBL(d,[
 ['close = O (weak close) ★','+0.31','+0.33','+0.30','dominant axis (~80% of suffix edge)'],
 ['EO (escape + O)','+0.53','+0.62','+0.61','capitulation reversal'],
 ['R2L (RSI2 oversold) ★','+0.31','+0.35','+0.30','same "buy-weakness" theme as close=O'],
 ['PB + R2L','+0.73','+0.50','+0.49','strongest line-5 combo'],
 ['gap = G3','+0.49','+0.50','+0.40','gap-up; G3-N best (+0.46/+0.62/+0.50)'],
 ['gap = G2','+0.07','+0.33','+0.30','gap-up (micro strong)'],
 ['wick = D (lower wick)','+0.10','+0.13','+0.18','bought the dip'],
 ['body = M (medium)','+0.03','+0.09','+0.09','> X (large, already extended)'],
 ['vol = B (controlled)','best','best','best','B>N>L>W>VB'],
 ['suffix EBO / NDO / …O','top','top','top','O-ending wins'],
],['component','SP','NQ','R2k','note'])

H(d,'BEAR profile — mirror (Z-signals)')
TBL(d,[
 ['close = A (strong close) ★','+0.08','best','best','mirror of bull O'],
 ['ne = E (escape)','+0.03','+0.08','+0.10','—'],
 ['wick = U (upper)','+0.02','+0.11','+0.06','—'],
 ['gap = G3/G2','+0.20','+0.40','+0.22','mean-reversion up even on Z'],
 ['suffix EBA / NHA / …A','top','top','top','A-ending wins on bear'],
],['component','SP','NQ','R2k','note'])

H(d,'✗ AVOID — universal losers / traps')
for x in ['VB volume (climactic) — WORST bucket over 5 years (−0.36 cross-uni); the v5 "VB is king" was a 2025 artifact.',
 'close=A on a BULL signal / close=O on a BEAR signal (wrong side of the dominant axis).',
 'range V/C WITHOUT a gap — micro trap (nasdaq −1.10); with a gap it becomes a breakout.',
 'PS + R2H (PSAR-sell + overbought) — worst line-5 combo (−0.5…−0.8).',
 'EUR & NH suffix (universally weak); ACC_TR as the preceding bar (−2.7 in micro).',
 'Russell2k RAW signals (all net-negative over 5 years) — need a composite/atomic filter.']:
    P(d,'•  '+x,sz=9)

H(d,'⚠ REGIME — the meta-rule')
P(d,'75% of composites and ~57% of 4-bar sequences are 2025-ARTIFACT (positive only in the recent risk-on '
  'regime). Only the STABLE set + the weak-close gap-up edge survive across years. The app already computes a '
  'RISK_ON/NEUTRAL/RISK_OFF label (ai_journal/regime.py); it currently reads RISK_OFF (as of 2026-06-09) → size down.',True,9)

H(d,'One-line recipe')
P(d,'LONG: bull T-signal · close=O · gap G2/G3 · (bonus: R2L oversold, vol=B, wick=D) · −15%/+100% exit · '
  'small size · stand down in RISK_OFF.  SHORT/avoid mirror: bear Z-signal · close=A.',True,10)
d.save(OUT+'/_MASTER_CHEATSHEET.docx')

# md twin
with open(OUT+'/_MASTER_CHEATSHEET.md','w') as f:
    f.write("""# TZ+WLNBB 5-Year Master Cheat-Sheet

## ★ WINNING EDGE — weak-close gap-up
**BULL T-signal AND close=O AND gap(G2/G3)** · entry next-open · −15%/+100% · $500k floor · small size.
EXPECT: sp500 +0.84 (n19k, win55) · r2k +0.70 (EO+gap) · nasdaq +0.51. Positive 5/6 years (only 2022 bear −).

## BULL atomic components (lift med10 SP/NQ/R2k)
- close=O **+0.31/+0.33/+0.30** (dominant) · EO +0.53/+0.62/+0.61 · R2L **+0.31/+0.35/+0.30** · PB+R2L +0.73/+0.50/+0.49
- gap G3 +0.49/+0.50/+0.40 (G3-N best) · G2 +0.07/+0.33/+0.30 · wick=D +0.10/+0.13/+0.18 · body=M · vol=B · suffix …O

## BEAR mirror (Z): close=A best · ne=E · gap G3/G2 mean-revert · suffix …A

## ✗ AVOID: VB volume (worst, v5 artifact) · wrong close-axis · range V/C no-gap (micro trap) · PS+R2H · EUR/NH suffix · ACC_TR prev · r2k raw signals

## ⚠ REGIME: 75% of composites are 2025-artifacts. App label = ai_journal/regime.py (now RISK_OFF). Size down in weak tape.

## ONE-LINE: LONG bull-T · close=O · gap G2/G3 · (+R2L/B/D) · −15%/+100% · small size · stand down RISK_OFF.
""")
print("cheat-sheet saved (docx + md) ->",OUT)
